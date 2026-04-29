from __future__ import annotations

from pathlib import Path
from typing import Set, Dict, Any
from datetime import datetime

from docxtpl import DocxTemplate
from docx import Document


BASE_DIR = Path(__file__).resolve().parent
VORLAGEN_DIR = BASE_DIR

OUTPUT_DIR = BASE_DIR / "Output_wordvorlage"
OUTPUT_DIR.mkdir(exist_ok=True)


def safe_filename(s: str) -> str:
    return "".join(c for c in (s or "").strip() if c.isalnum() or c in ("-", "_"))


def get_template_vars(tpl_name: str) -> Set[str]:
    tpl_path = VORLAGEN_DIR / tpl_name
    if not tpl_path.exists():
        raise FileNotFoundError(
            f"Vorlage nicht gefunden: {tpl_path}\n"
            f"Vorhandene .docx im Repo: {[p.name for p in VORLAGEN_DIR.glob('*.docx')]}"
        )

    tpl = DocxTemplate(str(tpl_path))
    return set(tpl.get_undeclared_template_variables() or [])


def _cell_text(cell) -> str:
    parts = []
    for p in cell.paragraphs:
        txt = p.text.replace("\xa0", " ").strip()
        if txt:
            parts.append(txt)
    return " ".join(parts).strip()


def _row_is_empty(row) -> bool:
    texts = [_cell_text(cell) for cell in row.cells]
    return all(not t for t in texts)


def _delete_row(row) -> None:
    tr = row._tr
    tr.getparent().remove(tr)


def _is_empty_or_zero(value: Any) -> bool:
    v = str(value or "").strip()
    return v in {"", "0,00", "0,00 €", "0.00", "0.00 €", "-", "--"}


def _cleanup_empty_table_rows(doc: Document) -> None:
    for table in doc.tables:
        rows = list(table.rows)

        for row in reversed(rows):
            if _row_is_empty(row):
                _delete_row(row)


def _row_contains_any_marker(row, markers: list[str]) -> bool:
    row_text = " ".join(_cell_text(cell) for cell in row.cells)
    return any(marker in row_text for marker in markers)


def _cleanup_optional_cost_rows(doc: Document, context: Dict[str, Any]) -> None:
    rows_to_match: list[list[str]] = []

    if _is_empty_or_zero(context.get("ABMELDEKOSTEN", "")):
        rows_to_match.append([
            "ABMELDEKOSTEN",
            "{{MELDUNGSKOSTEN}}",
            "Abmeldekosten",
        ])

    if _is_empty_or_zero(context.get("UMMELDEKOSTEN", "")):
        rows_to_match.append([
            "UMMELDEKOSTEN",
            "{{MELDUNGSKOSTEN}}",
            "Ummeldekosten",
        ])

    if _is_empty_or_zero(context.get("MELDUNGSKOSTEN", "")):
        rows_to_match.append([
            "MELDUNGSKOSTEN",
            "{{MELDUNGSKOSTEN}}",
            "Meldungskosten",
            "An- und Abmeldekosten",
            "An- & Abmeldekosten",
        ])

    if _is_empty_or_zero(context.get("ZUSATZKOSTEN_BEZEICHNUNG1", "")) or _is_empty_or_zero(context.get("ZUSATZKOSTEN_BETRAG1", "")):
        rows_to_match.append([
            "ZUSATZKOSTEN_BEZEICHNUNG1",
            "ZUSATZKOSTEN_BETRAG1",
            "{{ZUSATZKOSTEN_BEZEICHNUNG1}}",
            "{{ZUSATZKOSTEN_BETRAG1}}",
        ])

    if _is_empty_or_zero(context.get("ZUSATZKOSTEN_BEZEICHNUNG2", "")) or _is_empty_or_zero(context.get("ZUSATZKOSTEN_BETRAG2", "")):
        rows_to_match.append([
            "ZUSATZKOSTEN_BEZEICHNUNG2",
            "ZUSATZKOSTEN_BETRAG2",
            "{{ZUSATZKOSTEN_BEZEICHNUNG2}}",
            "{{ZUSATZKOSTEN_BETRAG2}}",
        ])

    if _is_empty_or_zero(context.get("ZUSATZKOSTEN_BEZEICHNUNG3", "")) or _is_empty_or_zero(context.get("ZUSATZKOSTEN_BETRAG3", "")):
        rows_to_match.append([
            "ZUSATZKOSTEN_BEZEICHNUNG3",
            "ZUSATZKOSTEN_BETRAG3",
            "{{ZUSATZKOSTEN_BEZEICHNUNG3}}",
            "{{ZUSATZKOSTEN_BETRAG3}}",
        ])

    for table in doc.tables:
        rows = list(table.rows)
        rows_to_delete = []

        for row in rows:
            for markers in rows_to_match:
                if _row_contains_any_marker(row, markers):
                    rows_to_delete.append(row)
                    break

        for row in reversed(rows_to_delete):
            _delete_row(row)


def render_word(tpl_name: str, context: Dict[str, Any], out_prefix: str) -> Path:
    tpl_path = VORLAGEN_DIR / tpl_name
    if not tpl_path.exists():
        raise FileNotFoundError(
            f"Vorlage nicht gefunden: {tpl_path}\n"
            f"Vorhandene .docx im Repo: {[p.name for p in VORLAGEN_DIR.glob('*.docx')]}"
        )

    clean_context = {k: ("" if v is None else v) for k, v in context.items()}

    tpl = DocxTemplate(str(tpl_path))
    tpl.render(clean_context)

    nachname = safe_filename(str(clean_context.get("MANDANT_NACHNAME", "Unbekannt") or "Unbekannt"))
    timestamp = datetime.now().strftime("%d-%m-%Y")
    out_name = f"01-AS_an_VR_{timestamp}.docx"
    out_path = OUTPUT_DIR / out_name

    # Erst rendern und speichern
    tpl.save(str(out_path))

    # Danach Word-Datei erneut öffnen und leere / optionale Tabellenzeilen löschen
    doc = Document(str(out_path))
    _cleanup_optional_cost_rows(doc, clean_context)
    _cleanup_empty_table_rows(doc)
    doc.save(str(out_path))

    return out_path
